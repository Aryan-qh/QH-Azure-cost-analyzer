"""
Word Document Generation Service
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
from typing import Dict, List
import os


class DocumentGeneratorService:
    """Generate Word documents for cost reports"""
    
    def __init__(self, output_directory: str):
        self.output_directory = output_directory
        os.makedirs(output_directory, exist_ok=True)
    
    def add_table_to_doc(self, doc: Document, table_data: List[list], headers: List[str], title: str = None):
        """Add a formatted table to the Word document"""
        
        if title:
            para = doc.add_paragraph()
            run = para.add_run(title)
            run.bold = True
            run.font.size = Pt(11)
        
        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data rows
        for row_data in table_data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add spacing
    
    def generate_cost_report(self, all_data: Dict, num_days: int) -> str:
        """Generate a Word document with cost data"""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Azure Cost Summary Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create date range string
        end_date = datetime.now() - timedelta(days=1)
        start_date = end_date - timedelta(days=num_days - 1)
        
        # Format dates with day names
        date_list = []
        for i in range(num_days):
            date = start_date + timedelta(days=i)
            date_list.append(f"{date.strftime('%A')} ({date.strftime('%m/%d')})")
        
        date_range_str = ", ".join(date_list)
        
        # Add greeting
        greeting = doc.add_paragraph()
        greeting.add_run("Hi Team,\n\n").bold = False
        greeting.add_run(
            f"Please find below the Azure cost summary for {date_range_str} "
            f"for all subscriptions, along with percentage changes compared to the previous day.\n"
        )
        
        # Add tables for each subscription
        for sub_name in ['prod', 'dev', 'test', 'main']:
            if sub_name in all_data and all_data[sub_name]:
                data = all_data[sub_name]
                
                # Add subscription header
                doc.add_heading(f'{sub_name.capitalize()} Environment', level=2)
                
                # Add cost table
                self.add_table_to_doc(doc, data['cost_table'], data['headers'])
                
                # Add percentage difference table
                self.add_table_to_doc(
                    doc, 
                    data['percent_table'], 
                    data['headers'],
                    f"Percentage difference for {sub_name}"
                )
        
        # Add closing
        doc.add_paragraph("\nThank you.")
        
        # Save document
        filename = f"Azure_Cost_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_directory, filename)
        doc.save(filepath)
        
        return filename
    
    def prepare_report_data(
        self,
        subscription_id: str,
        subscription_name: str,
        num_days: int,
        cost_data_service,
        cost_processor
    ) -> Dict:
        """Prepare data for a subscription report"""
        
        # Calculate date range
        end_date = datetime.now() - timedelta(days=1)
        start_date = end_date - timedelta(days=num_days - 1)
        
        # Get all data in one API call
        response_data = cost_data_service.get_cost_data_range(
            subscription_id, start_date, end_date
        )
        
        if not response_data:
            return None
        
        daily_data = cost_data_service.parse_range_response(response_data)
        
        # Prepare data structures
        cost_table_data = []
        percent_table_data = []
        all_costs = []
        date_strings = []
        
        # Process each day
        for i in range(num_days - 1, -1, -1):
            date = datetime.now() - timedelta(days=i + 1)
            date_key = int(date.strftime('%Y%m%d'))
            date_str = date.strftime('%m/%d')
            date_strings.append(date_str)
            
            day_rows = daily_data.get(date_key, [])
            costs = cost_processor.process_cost_data(day_rows)
            all_costs.append(costs)
        
        # Determine categories
        categories = cost_processor.get_relevant_categories(all_costs, subscription_name)
        
        # Build cost table
        for i, costs in enumerate(all_costs):
            row = [date_strings[i]]
            for category in categories:
                row.append(f"${costs[category]:.2f}")
            cost_table_data.append(row)
        
        # Build percentage change table
        for i in range(1, len(all_costs)):
            row = [date_strings[i]]
            
            for category in categories:
                prev_cost = all_costs[i - 1][category]
                curr_cost = all_costs[i][category]
                
                percent_change = cost_processor.calculate_percentage_change(
                    prev_cost, curr_cost
                )
                
                row.append(f"{percent_change:+.2f}%")
            
            percent_table_data.append(row)
        
        headers = ['Date'] + categories
        
        return {
            'cost_table': cost_table_data,
            'percent_table': percent_table_data,
            'headers': headers,
            'date_strings': date_strings
        }